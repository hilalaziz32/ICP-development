"""
build_brief_docx.py — Render the structured 8-field brief into a .docx with
the standard Scaletopia layout.

USAGE:
    python build_brief_docx.py --input brief.json --output "2025-11-15-DTC-VP-of-Marketing-brief.docx"

INPUT FORMAT (brief.json):
    {
        "metadata": {
            "client": "Kinship",
            "segment": "DTC Supplements",
            "persona": "VP of Marketing",
            "date_completed": "2025-11-15",
            "completed_by": "sms-brief skill v1"
        },
        "fields": {
            "1_buyer": "Director of Marketing at DTC supplement brands... [source: ...]",
            "2_pains": [
                {
                    "text": "Blended CAC jumped from $89 to $134...",
                    "score": 12,
                    "verification": "verified",
                    "sources": ["transcript-005.docx, ~18:45", "Master Sheet Tab 2"]
                },
                ...
            ],
            "3_their_language": {
                "problem_quotes": [{"quote": "...", "source": "..."}, ...],
                "success_quotes": [...],
                "industry_terms": [...],
                "outsider_phrases": [...]
            },
            "4a_industry_beliefs": [{"belief": "...", "why_wrong": "...", "source": "..."}, ...],
            "4b_service_objections": [{"objection": "...", "dissolving_mechanism": "...", "source": "..."}, ...],
            "5_dream_outcome": [{"text": "...", "source": "..."}, ...],
            "6_sophistication": {"level": "Mid", "evidence": "...", "source": "..."},
            "7_reply_behaviour": {"replies_favorably_to": [...], "ignores": [...], "gap": "..."}
        },
        "source_inventory": { ... output of source_inventory.py ... },
        "qa_report": {
            "fields_checked": 8,
            "claims_passed": 14,
            "flagged_unverified": 2,
            "rewrites_made": ["Field 2 #3 rewritten from generic..."],
            "drops_made": ["Field 3 dropped fake-verbatim quote..."],
            "gaps_present": ["Field 7"]
        }
    }

DEPENDENCIES:
    pip install python-docx --break-system-packages

OUTPUT:
    A .docx file with:
    - Title page (client + segment + persona + date)
    - 8 fields as proper sections
    - Inline source citations
    - Appendix: source inventory + QA report
"""

import argparse
import json
import sys


def _check_python_docx():
    try:
        import docx  # noqa: F401
    except ImportError:
        print(
            "ERROR: python-docx not installed. Run:\n"
            "    pip install python-docx --break-system-packages",
            file=sys.stderr
        )
        sys.exit(1)


def render_brief(brief: dict, output_path: str):
    _check_python_docx()
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    md = brief.get("metadata", {})
    fields = brief.get("fields", {})

    # === Title ===
    title = doc.add_heading(
        f"{md.get('client', '?')} — Layer A Targeting Brief",
        level=0
    )
    doc.add_paragraph(f"Segment: {md.get('segment', '?')}")
    doc.add_paragraph(f"Persona: {md.get('persona', '?')}")
    doc.add_paragraph(f"Date: {md.get('date_completed', '?')}")
    doc.add_paragraph(f"Completed by: {md.get('completed_by', '?')}")
    doc.add_page_break()

    # === Field 1: Buyer ===
    doc.add_heading("1. Buyer", level=1)
    doc.add_paragraph(fields.get("1_buyer", "[GAP — no source material]"))

    # === Field 2: Top pains ===
    doc.add_heading("2. Top Pains (Ranked by Source Frequency)", level=1)
    for i, pain in enumerate(fields.get("2_pains", []), start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{pain.get('text', '')}").bold = True
        if pain.get("verification") != "verified":
            run = p.add_run(f"  [{pain.get('verification', '')}]")
            run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)  # orange-ish
        for src in pain.get("sources", []):
            doc.add_paragraph(f"  source: {src}", style="Intense Quote")
        doc.add_paragraph(f"  score: {pain.get('score', '?')}")

    # === Field 3: Their language ===
    doc.add_heading("3. Their Language (Verbatim)", level=1)
    lang = fields.get("3_their_language", {})

    doc.add_heading("How they describe the problem", level=2)
    for q in lang.get("problem_quotes", []):
        doc.add_paragraph(f'"{q.get("quote", "")}"  — {q.get("source", "")}', style="Quote")

    doc.add_heading("How they describe success", level=2)
    for q in lang.get("success_quotes", []):
        doc.add_paragraph(f'"{q.get("quote", "")}"  — {q.get("source", "")}', style="Quote")

    doc.add_heading("Industry terms to weave in", level=2)
    for term in lang.get("industry_terms", []):
        doc.add_paragraph(f"- {term}", style="List Bullet")

    doc.add_heading("Outsider phrases — AVOID", level=2)
    for phrase in lang.get("outsider_phrases", []):
        doc.add_paragraph(f"- {phrase}", style="List Bullet")

    # === Field 4a: Industry beliefs ===
    doc.add_heading("4a. Industry Mistaken Beliefs", level=1)
    for b in fields.get("4a_industry_beliefs", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"\"{b.get('belief', '')}\"").italic = True
        doc.add_paragraph(f"  Why this is wrong: {b.get('why_wrong', '')}")
        doc.add_paragraph(f"  source: {b.get('source', '')}", style="Intense Quote")

    # === Field 4b: Service objections ===
    doc.add_heading("4b. Service Hidden Objections (what blocks replies)", level=1)
    for o in fields.get("4b_service_objections", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"\"{o.get('objection', '')}\"").italic = True
        doc.add_paragraph(f"  Mechanism that dissolves it: {o.get('dissolving_mechanism', '')}")
        doc.add_paragraph(f"  source: {o.get('source', '')}", style="Intense Quote")

    # === Field 5: Dream outcome ===
    doc.add_heading("5. Dream Outcome (in their words)", level=1)
    for d in fields.get("5_dream_outcome", []):
        doc.add_paragraph(f'"{d.get("text", "")}"  — {d.get("source", "")}', style="Quote")

    # === Field 6: Sophistication ===
    doc.add_heading("6. Sophistication Level", level=1)
    soph = fields.get("6_sophistication", {})
    p = doc.add_paragraph()
    p.add_run(f"Level: {soph.get('level', '?')}").bold = True
    doc.add_paragraph(f"Evidence: {soph.get('evidence', '')}")
    doc.add_paragraph(f"source: {soph.get('source', '')}", style="Intense Quote")

    # === Field 7: Reply behaviour ===
    doc.add_heading("7. Reply Behaviour", level=1)
    rb = fields.get("7_reply_behaviour", {})
    if rb.get("gap"):
        p = doc.add_paragraph(f"GAP: {rb['gap']}")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    else:
        doc.add_heading("Replies favourably to:", level=2)
        for item in rb.get("replies_favorably_to", []):
            doc.add_paragraph(f"- {item}", style="List Bullet")
        doc.add_heading("Ignores / negative-replies to:", level=2)
        for item in rb.get("ignores", []):
            doc.add_paragraph(f"- {item}", style="List Bullet")

    # === Appendix: Source Inventory ===
    doc.add_page_break()
    doc.add_heading("Appendix — Source Inventory", level=1)
    inv = brief.get("source_inventory", {})
    if "summary" in inv:
        s = inv["summary"]
        doc.add_paragraph(f"Transcripts: {s.get('transcript_count', 0)}")
        doc.add_paragraph(f"Master Sheet present: {s.get('master_sheet_present', False)}")
        doc.add_paragraph(f"Onboarding form present: {s.get('onboarding_present', False)}")
        doc.add_paragraph(f"Case studies present: {s.get('case_studies_present', False)}")
        doc.add_paragraph(f"Prior briefs: {s.get('prior_briefs_count', 0)}")
    if inv.get("gaps"):
        doc.add_heading("Gaps", level=2)
        for g in inv["gaps"]:
            doc.add_paragraph(f"- {g}", style="List Bullet")

    # === Appendix: QA Report ===
    doc.add_heading("Appendix — QA Report", level=1)
    qa = brief.get("qa_report", {})
    doc.add_paragraph(f"Fields checked: {qa.get('fields_checked', 0)}")
    doc.add_paragraph(f"Claims passed: {qa.get('claims_passed', 0)}")
    doc.add_paragraph(f"Flagged as unverified: {qa.get('flagged_unverified', 0)}")
    if qa.get("rewrites_made"):
        doc.add_heading("Rewrites this pass", level=2)
        for r in qa["rewrites_made"]:
            doc.add_paragraph(f"- {r}", style="List Bullet")
    if qa.get("drops_made"):
        doc.add_heading("Drops this pass", level=2)
        for d in qa["drops_made"]:
            doc.add_paragraph(f"- {d}", style="List Bullet")
    if qa.get("gaps_present"):
        doc.add_heading("Fields GAP-marked", level=2)
        for g in qa["gaps_present"]:
            doc.add_paragraph(f"- {g}", style="List Bullet")

    doc.save(output_path)
    print(f"Wrote brief to {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Render structured brief JSON to .docx")
    parser.add_argument("--input", required=True, help="Path to brief JSON")
    parser.add_argument("--output", required=True, help="Path to write .docx")
    args = parser.parse_args()

    with open(args.input) as f:
        brief = json.load(f)

    render_brief(brief, args.output)


if __name__ == "__main__":
    main()
